# Test should be run inside the base folder of the repository: i.e. `PeerReviweWordCounter` folder
from docx import Document
from docx.shared import RGBColor

def create_test_word_file(file_name):
    doc = Document()

    p = doc.add_paragraph()
    r = p.add_run("This is red text. ")
    r.font.color.rgb = RGBColor(255, 0, 0)

    r = p.add_run("This is blue text. Let us add some more blue \n\n\n ")
    r.font.color.rgb = RGBColor(0, 0, 255)

    r = p.add_run("This is green text. Greeny Green ")
    r.font.color.rgb = RGBColor(0, 255, 0)

    doc.save(file_name)

create_test_word_file("test_doc.docx")


import unittest
from word_count import extract_color_and_word_count

class TestWordCounter(unittest.TestCase):

    def test_multicolor_text(self):
        result = extract_color_and_word_count('test_doc.docx')
        expected_result = {'red': 4, 'blue': 10, 'lime': 6}  # Adjusted color names and cases.
        self.assertEqual(result, expected_result)


if __name__ == '__main__':
    unittest.main()

